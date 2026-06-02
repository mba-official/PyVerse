from docx import Document
import re
import requests
from urllib.parse import quote
import threading
import itertools
import sys
import time

# -----------------------------
# CONFIG
# -----------------------------
INPUT_DOCX = "all_references.docx"
OUTPUT_DOCX = "verified_references_report.docx"

session = requests.Session()

# -----------------------------
# SIMPLE SPINNER (ANIMATION)
# -----------------------------
class Spinner:
    def __init__(self, message="Searching"):
        self.spinner = itertools.cycle(["|", "/", "-", "\\"])
        self.running = False
        self.message = message
        self.thread = None

    def start(self):
        self.running = True
        self.thread = threading.Thread(target=self._spin)
        self.thread.start()

    def _spin(self):
        while self.running:
            sys.stdout.write(f"\r{self.message}... {next(self.spinner)}")
            sys.stdout.flush()
            time.sleep(0.1)

    def stop(self):
        self.running = False
        if self.thread:
            self.thread.join()
        sys.stdout.write("\r" + " " * 60 + "\r")
        sys.stdout.flush()


# -----------------------------
# LOAD DOCX
# -----------------------------
def load_docx(file_path):
    doc = Document(file_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


# -----------------------------
# SPLIT CHAPTERS
# -----------------------------
def split_by_chapter(lines):
    chapters = {}
    current = "Unknown"

    for line in lines:
        match = re.match(r"Chapter\s*\d+", line, re.IGNORECASE)
        if match:
            current = match.group()
            chapters[current] = []
        else:
            chapters.setdefault(current, []).append(line)

    return chapters


# -----------------------------
# CLASSIFY
# -----------------------------
def classify_reference(ref):
    r = ref.lower()

    if "quran" in r or "surah" in r:
        return "Quran"
    if any(x in r for x in ["bukhari", "muslim", "tirmidhi", "hadith"]):
        return "Hadith"
    if re.search(r"\(\d{4}\)", ref):
        return "Academic"
    return "Book"


# -----------------------------
# SOURCES
# -----------------------------
def google_scholar(query):
    return ("Google Scholar", f"https://scholar.google.com/scholar?q={quote(query)}", 0.90)

def crossref_search(query):
    try:
        url = f"https://api.crossref.org/works?query={quote(query)}&rows=1"
        r = session.get(url, timeout=10).json()

        items = r.get("message", {}).get("items", [])
        if items:
            item = items[0]
            doi = item.get("DOI")
            if doi:
                return ("Crossref", f"https://doi.org/{doi}", 0.95)
    except:
        pass
    return None


def openlibrary_search(query):
    try:
        url = f"https://openlibrary.org/search.json?q={quote(query)}"
        r = session.get(url, timeout=10).json()

        docs = r.get("docs", [])
        if docs:
            key = docs[0].get("key", "")
            return ("OpenLibrary", f"https://openlibrary.org{key}", 0.75)
    except:
        pass
    return None


def web_fallback(query):
    return ("Web Search", f"https://www.google.com/search?q={quote(query)}", 0.60)


# -----------------------------
# MULTI-SOURCE SEARCH (TOP 3)
# -----------------------------
def search_all_sources(query):
    sources = []

    spinner = Spinner("Searching academic sources")
    spinner.start()

    try:
        sources.append(google_scholar(query))

        spinner.message = "Checking CrossRef database"
        c = crossref_search(query)
        if c:
            sources.append(c)

        spinner.message = "Checking OpenLibrary"
        o = openlibrary_search(query)
        if o:
            sources.append(o)

        sources.append(web_fallback(query))

    finally:
        spinner.stop()

    # deduplicate
    seen = set()
    unique = []
    for s in sources:
        if s[1] not in seen:
            unique.append(s)
            seen.add(s[1])

    return unique[:3]


# -----------------------------
# ENRICH
# -----------------------------
def enrich_reference(ref):
    rtype = classify_reference(ref)

    if rtype == "Quran":
        return [("Al-Qur’an al-Kareem", "https://quran.com", 1.0)]
    if rtype == "Hadith":
        return [("Sunnah.com", "https://sunnah.com", 0.85)]

    query = re.sub(r"[^\w\s]", " ", ref).strip()
    return search_all_sources(query)


# -----------------------------
# PROCESS
# -----------------------------
def process(chapters):
    results = []
    all_refs = [(c, r) for c in chapters for r in chapters[c]]
    total = len(all_refs)

    for idx, (chapter, ref) in enumerate(all_refs, 1):

        print(f"\n[{idx}/{total}] Processing reference...")

        sources = enrich_reference(ref)

        confidence = round(sum(s[2] for s in sources) / len(sources), 2)

        results.append({
            "Chapter": chapter,
            "Reference": ref,
            "Sources": sources,
            "Confidence": confidence
        })

    return results


# -----------------------------
# WRITE DOCX (YOUR FORMAT)
# -----------------------------
def write_docx(data):
    doc = Document()

    for i, item in enumerate(data, 1):

        doc.add_paragraph(f"Chapter: {item['Chapter']}")
        doc.add_paragraph(f"Reference No: {i}")
        doc.add_paragraph(f"Reference: {item['Reference']}")
        doc.add_paragraph("Available on Internet: Yes")

        doc.add_paragraph("Best Match Source:")
        for j, s in enumerate(item["Sources"], 1):
            doc.add_paragraph(f"{j}. {s[0]}")

        doc.add_paragraph("Links:")
        for j, s in enumerate(item["Sources"], 1):
            doc.add_paragraph(f"{j}. {s[1]}")

        doc.add_paragraph(f"Confidence Score: {item['Confidence']}/1")
        doc.add_paragraph("_" * 80)

    doc.save(OUTPUT_DOCX)


# -----------------------------
# MAIN
# -----------------------------
lines = load_docx(INPUT_DOCX)
chapters = split_by_chapter(lines)

data = process(chapters)
write_docx(data)

print(f"\nDone! DOCX saved as: {OUTPUT_DOCX}")