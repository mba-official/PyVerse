from docx import Document
from rapidfuzz import fuzz
import pandas as pd
import re
import unicodedata

# =========================
# CONFIG
# =========================
INPUT_FILE = "references.docx"
OUTPUT_DOCX = "clean_bibliography.docx"
OUTPUT_XLSX = "merge_report.xlsx"

FUZZY_THRESHOLD = 90


# =========================
# TEXT NORMALIZATION
# =========================
def normalize_text(text):
    text = text.lower()

    # remove diacritics (important for arabic/transliterations)
    text = unicodedata.normalize('NFKD', text)
    text = ''.join([c for c in text if not unicodedata.combining(c)])

    text = re.sub(r'[^a-z0-9 ]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    return text


# =========================
# CLEANING FUNCTION
# =========================
def clean_entry(text):
    text = text.strip()

    # -------------------------
    # QURAN RULE (merge all)
    # -------------------------
    if text.lower().startswith("surah") or "quran" in text.lower():
        return "QURAN"

    # -------------------------
    # REMOVE PAGE NUMBERS
    # -------------------------
    text = re.sub(r'\b(pg|pp|p)\.?\s*\d+(\s*[-–]\s*\d+)?', '', text, flags=re.I)

    # -------------------------
    # REMOVE VOLUMES
    # -------------------------
    text = re.sub(r'vol\.?\s*\d+', '', text, flags=re.I)

    # -------------------------
    # REMOVE EXTRA COMMAS/BRACKETS SPACING
    # -------------------------
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s+,', ',', text)

    return text.strip()


# =========================
# CANONICAL KEY BUILDER
# (VERY IMPORTANT FOR MERGING)
# =========================
def build_key(text):
    n = normalize_text(text)

    # ---- ISLAMIC CORE TEXTS ----
    if "ghazali" in n and "ihya" in n:
        return "ghazali_ihya"
    if "tabari" in n:
        return "tabari_tafsir"
    if "ibn kathir" in n:
        return "ibn_kathir_tafsir"
    if "qurtubi" in n:
        return "qurtubi_tafsir"
    if "razi" in n:
        return "razi_tafsir"

    # ---- HADITH COLLECTIONS ----
    if "bukhari" in n:
        return "sahih_bukhari"
    if "muslim" in n:
        return "sahih_muslim"
    if "tirmidhi" in n:
        return "jami_tirmidhi"

    # ---- PSYCHOLOGY CLASSICS ----
    if "goleman" in n:
        return "goleman_emotional_intelligence"
    if "frankl" in n:
        return "frankl_meaning"
    if "kahneman" in n:
        return "kahneman_thinking_fast_slow"
    if "bandura" in n:
        return "bandura_social_learning"

    # fallback: fuzzy-safe key
    return n[:60]


# =========================
# READ FILE
# =========================
doc = Document(INPUT_FILE)

raw = []
quran_found = False

for p in doc.paragraphs:
    if p.text.strip():
        cleaned = clean_entry(p.text)

        if cleaned == "QURAN":
            quran_found = True
        elif cleaned:
            raw.append(cleaned)


# =========================
# GROUPING + DEDUPLICATION
# =========================
groups = {}
merge_log = []

for item in raw:
    key = build_key(item)

    if key not in groups:
        groups[key] = item
    else:
        # fuzzy check for near duplicates
        existing = groups[key]
        score = fuzz.ratio(normalize_text(item), normalize_text(existing))

        if score > FUZZY_THRESHOLD:
            merge_log.append({
                "original": item,
                "merged_into": existing
            })
        else:
            # keep longer/more complete version
            if len(item) > len(existing):
                merge_log.append({
                    "original": existing,
                    "merged_into": item
                })
                groups[key] = item
            else:
                merge_log.append({
                    "original": item,
                    "merged_into": existing
                })


# =========================
# FINAL LIST
# =========================
final_list = list(groups.values())

if quran_found:
    final_list.append("Al-Qur’an al-Karim")


# sort alphabetically
final_list.sort(key=lambda x: x.lower())


# =========================
# WRITE DOCX
# =========================
out = Document()
out.add_heading("Bibliography (Cleaned)", level=1)

for r in final_list:
    out.add_paragraph(r)

out.save(OUTPUT_DOCX)


# =========================
# WRITE REPORT (IMPORTANT FOR PHD)
# =========================
df = pd.DataFrame(merge_log)
df.to_excel(OUTPUT_XLSX, index=False)


# =========================
# SUMMARY
# =========================
print("\nDONE ✅")
print("Original entries:", len(raw))
print("Final unique entries:", len(final_list))
print("Saved:", OUTPUT_DOCX)
print("Merge log:", OUTPUT_XLSX)